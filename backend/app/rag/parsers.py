from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Protocol

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.config import settings
from app.rag.images import ExtractedImage, usable_image
from app.rag.prompts import TRANSCRIBE_PROMPT
from app.rag.vision import VisionProvider


@dataclass
class ParsedDocument:
    text: str
    images: list[ExtractedImage]


class DocumentParser(Protocol):
    def parse(self, data: bytes) -> ParsedDocument: ...


class TextParser:
    def parse(self, data: bytes) -> ParsedDocument:
        return ParsedDocument(text=data.decode("utf-8"), images=[])


class DocxParser:
    def parse(self, data: bytes) -> ParsedDocument:
        document = DocxDocument(BytesIO(data))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        return ParsedDocument(text=text, images=self.extract_images(document))

    def extract_images(self, document) -> list[ExtractedImage]:
        try:
            blobs = [
                part.blob
                for part in document.part.package.iter_parts()
                if part.content_type.startswith("image/")
                and "/media/" in str(part.partname)
            ]
        except Exception:
            return []
        images = (usable_image(blob) for blob in blobs)
        return [image for image in images if image is not None]


class ImageParser:
    def __init__(self, vision: VisionProvider) -> None:
        self.vision = vision

    def parse(self, data: bytes) -> ParsedDocument:
        text = self.vision.describe(data, TRANSCRIBE_PROMPT)
        original = usable_image(data)
        return ParsedDocument(text=text, images=[original] if original else [])


class PdfParser:
    def __init__(self, vision: VisionProvider) -> None:
        self.vision = vision

    def parse(self, data: bytes) -> ParsedDocument:
        reader = PdfReader(BytesIO(data))
        text = "\n\n".join(page.extract_text() for page in reader.pages)
        if not text.strip():
            text = self.ocr(data)
        return ParsedDocument(text=text, images=self.extract_images(reader))

    def extract_images(self, reader: PdfReader) -> list[ExtractedImage]:
        embedded = (
            data for page in reader.pages for data in self.page_images(page)
        )
        images = (usable_image(data) for data in embedded)
        return [image for image in images if image is not None]

    def page_images(self, page) -> list[bytes]:
        try:
            return [image.data for image in page.images]
        except Exception:
            return []

    def ocr(self, data: bytes) -> str:
        import pypdfium2 as pdfium

        document = pdfium.PdfDocument(data)
        pages = []
        for index in range(min(len(document), settings.ocr_max_pages)):
            bitmap = document[index].render(scale=2)
            buffer = BytesIO()
            bitmap.to_pil().save(buffer, "PNG")
            pages.append(self.vision.describe(buffer.getvalue(), TRANSCRIBE_PROMPT))
        return "\n\n".join(pages)


def build_parsers(vision: VisionProvider) -> dict[str, DocumentParser]:
    text_parser = TextParser()
    image_parser = ImageParser(vision)
    return {
        ".txt": text_parser,
        ".md": text_parser,
        ".pdf": PdfParser(vision),
        ".docx": DocxParser(),
        ".png": image_parser,
        ".jpg": image_parser,
        ".jpeg": image_parser,
    }


def parser_for(
    filename: str, parsers: dict[str, DocumentParser]
) -> DocumentParser | None:
    return parsers.get(Path(filename).suffix.lower())


def supported_suffixes(parsers: dict[str, DocumentParser]) -> str:
    return ", ".join(sorted(parsers))
